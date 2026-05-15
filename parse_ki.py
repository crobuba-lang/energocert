#!/usr/bin/env python3
"""parse_ki.py – Fast KI Expert .docx parser using lxml for table 9"""
import sys, re, json, zipfile
from docx import Document
from lxml import etree

def sf(s):
    try: return float(str(s).replace(',','.').strip())
    except: return None

def uv(cells):
    seen=set(); r=[]
    for c in cells:
        v=c.strip()
        if v and v!='-' and v not in seen: seen.add(v); r.append(v)
    return r

def extract(path):
    data = {}
    doc = Document(path)

    # ── TABLE 1 ──────────────────────────────────────────────────────
    if len(doc.tables) > 1:
        for row in doc.tables[1].rows:
            cells = [c.text.strip() for c in row.cells]
            k = cells[0] if cells else ''
            v = cells[-1] if len(cells)>1 else ''
            if 'INVESTITOR' in k.upper() or 'Investitor' in k: data['narucitelj'] = v
            if 'Naziv zgrade' in k: data['gradjevina'] = v
            if 'Vrsta zgrade' in k:
                vr = v.strip()
                if 'obile' in vr.lower() or 'obiteljsk' in vr.lower():
                    data['vrsta'] = 'Obiteljska stambena zgrada'
                elif 'višestamb' in vr.lower() or 'visestamb' in vr.lower():
                    data['vrsta'] = 'Višestambena zgrada'
                elif 'uredsk' in vr.lower() or 'poslovn' in vr.lower():
                    data['vrsta'] = 'Uredska zgrada'
                elif 'hotel' in vr.lower(): data['vrsta'] = 'Hotel/restoran'
                elif 'bolnic' in vr.lower(): data['vrsta'] = 'Bolnica'
                elif 'sportsk' in vr.lower(): data['vrsta'] = 'Sportska dvorana'
                elif 'trgovin' in vr.lower(): data['vrsta'] = 'Zgrada trgovine'
                else: data['vrsta'] = vr
            if 'k.č' in k or 'k.o.' in k.lower(): data['katastar'] = v
            if 'Adresa' in k: data['lokacija'] = v
            if 'Oplošje' in k and '(' in k: data['oplosje'] = v
            if 'Obujam' in k and 'V e' in k: data['obujam'] = v
            if 'Faktor oblika' in k: data['faktor'] = v
            if 'korisne površine' in k: data['ak'] = v
            if 'Meteorološka' in k: data['meteo'] = v

    # ── TABLE 2 – QHnd ref, QCnd, Htr ────────────────────────────────
    if len(doc.tables) > 2:
        rows2 = list(doc.tables[2].rows)
        for i, row in enumerate(rows2):
            cells = [c.text.strip() for c in row.cells]
            u = uv(cells); k = u[0] if u else ''
            nums = [x for x in u if sf(x) is not None]
            if 'Q H,nd' in k and 'kWh/a' in k and 'jedinici' not in k:
                data['qhndKwh'] = u[-1] if len(u)>1 else ''
            if "Q'' H,nd" in k and 'kWh/(m 2' in k and len(nums)>=2:
                data['qhndMax'] = nums[0]; data['qhndM2'] = nums[-1]
            if 'Q C,nd' in k and 'kWh/a' in k and 'jedinici' not in k:
                data['qcndKwh'] = u[-1] if len(u)>1 else ''
            if "Q'' C,nd" in k and 'kWh/(m 2' in k and len(nums)>=2:
                data['qcndM2'] = nums[-1]
            if 'H tr,adj' in k and 'W/(m 2' in k and len(nums)>=2:
                data['htrMax'] = nums[0]; data['htrAdj'] = nums[-1]

    # ── TABLE 3 – QHnd spec if exists ─────────────────────────────────
    # KI Expert specific data (second file) has same structure
    # Will be filled from spec file if provided

    # ── TABLE 4 – Edel, Eprim, OIE ref ───────────────────────────────
    if len(doc.tables) > 4:
        for row in doc.tables[4].rows:
            cells = [c.text.strip() for c in row.cells]
            u = uv(cells); k = u[0] if u else ''
            if 'isporučena energija' in k and 'termotehnički' in k and len(u)>1:
                data['edel'] = u[1]
            if 'primarna energija' in k and 'termotehnički' in k and len(u)>1:
                data['eprim'] = u[1]
            for x in u:
                f = sf(x)
                if f and 1<f<100 and '%' in ' '.join(u): data['oieUdio'] = x

    # ── TABLE 5 – Eprim/m2, nZEB ─────────────────────────────────────
    if len(doc.tables) > 5:
        rows5 = list(doc.tables[5].rows)
        for i, row in enumerate(rows5):
            cells = [c.text.strip() for c in row.cells]
            u = uv(cells); k = u[0] if u else ''
            if 'E del' in k and 'kWh/a' in k and len(u)>1:
                data['edel'] = u[1]
            if 'E prim' in k and 'kWh/a' in k and 'jedinici' not in k and len(u)>1:
                data['eprim'] = u[1]
            if 'jedinici' in k and 'E prim' in k and i+1<len(rows5):
                nu = uv([c.text.strip() for c in rows5[i+1].cells])
                nums = [x for x in nu if sf(x) is not None]
                if len(nums)>=2: data['eprimMax']=nums[0]; data['eprimM2']=nums[-1]
            if 'nZEB' in ' '.join(u): data['nzeb'] = 'da'

    # ── TABLE 8 – grijanje, energent ─────────────────────────────────
    if len(doc.tables) > 8:
        for row in doc.tables[8].rows:
            cells = [c.text.strip() for c in row.cells]
            u = uv(cells); k = u[0] if u else ''
            if k == 'Sustav grijanja:' and len(u)>1: data['grijVrsta'] = u[1]
            if 'energenta za grijanje' in k and len(u)>1: data['grijEnergent'] = u[1]
            if 'obnovljive energije' in k and len(u)>1: data['oieUdio'] = u[1]

    # ── SEARCH ALL TABLES – koeficijenti transmisijskih gubitaka ─────
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            # Join all cell text for full row search
            full_row = ' '.join(cells)
            u = uv(cells)
            # Get all numeric values from the row
            all_nums = [c for c in cells if sf(c) is not None]
            if not all_nums:
                continue
            val = all_nums[-1]  # last number in row
            if 'prema vanjskom okolišu' in full_row and ('H D' in full_row or 'H_D' in full_row or 'HD' in full_row):
                data['hD'] = val
            if 'prema tlu' in full_row and ('g,avg' in full_row or 'Hg' in full_row or 'H g' in full_row):
                data['hGavg'] = val
            if 'kroz negrijani' in full_row and ('H U' in full_row or 'HU' in full_row):
                data['hU'] = val
            if 'prema susjednoj' in full_row and ('H A' in full_row or 'HA' in full_row):
                data['hA'] = val
            if ('Ukupni koeficijent' in full_row or 'Ukupni H' in full_row) and 'izmjene topline' in full_row:
                data['hTr'] = val

    # ── TABLE 9 – XML FAST parsing ────────────────────────────────────
    NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    W = '{' + NS + '}'

    def get_text(el):
        return ''.join(t.text or '' for t in el.iter(W+'t')).strip()

    with zipfile.ZipFile(path, 'r') as z:
        xml = z.read('word/document.xml')
    root = etree.fromstring(xml)
    all_tbls = root.findall('.//' + W + 'tbl')

    if len(all_tbls) > 9:
        t9_rows = []
        for tr in all_tbls[9].findall('.//' + W+'tr'):
            cells = [get_text(tc) for tc in tr.findall('.//' + W+'tc')]
            if cells: t9_rows.append(cells)

        uvals = []; seen_names = set(); raw_lines = []
        VALID = ['Vanjski zid','Zid prema','Pod na tlu','Strop','Kosi krov','Ravni krov']
        SKIP  = ['2.A.','Unutarnja','Opći','Toplinska','Površinska','Dinamičke',
                 'Slojevi','Ispravci','Proračun','Odabrani','Naziv otvora',
                 'Gibanj','Stacionarni','Granice','Nema','Koriste','Popis',
                 'Gubitak','A gd','Q sol','Q int','Q Ve','n inf','Δn','Godišnje',
                 'Svi mj','Siječanj','Veljača','Ožujak','Travanj','Svibanj',
                 'Lipanj','Srpanj','Kolovoz','Rujan','Listopad','Studeni',
                 'Prosinac','Električna','SPF','Obnovlj','Ukupni','Ako je',
                 'Smještaj','Tip ']

        for cells_text in t9_rows:
            # Compact unique for raw
            seen=set(); unique=[]
            for c in cells_text:
                if c and c not in seen: seen.add(c); unique.append(c)
            if len(raw_lines)<2000 and unique:
                raw_lines.append(' | '.join(unique[:6]))

            full = ' '.join(cells_text[:15])

            # n50
            if 'Broj izmjena zraka' in full and '50 Pa' in full:
                m = re.search(r'n\s*50\s*=\s*([\d.]+)', ' '.join(cells_text))
                if m: data['zrakN50'] = m.group(1)

            # Geometry
            if 'Oplošje grijanog dijela zgrade' in full and 'A,' in full:
                nums = list(dict.fromkeys([c for c in cells_text if re.match(r'^\d{2,}\.\d+$',c)]))
                if nums: data['oplosje'] = nums[-1]
            if 'Obujam grijanog dijela zgrade' in full and 'V e,' in full:
                nums = list(dict.fromkeys([c for c in cells_text if re.match(r'^\d{2,}\.\d+$',c)]))
                if nums: data['obujam'] = nums[-1]
            if 'Obujam grijanog zraka' in full:
                nums = list(dict.fromkeys([c for c in cells_text if re.match(r'^\d{2,}\.\d+$',c)]))
                if nums: data['obujamZrak'] = nums[-1]
            if 'Faktor oblika zgrade' in full:
                nums = list(dict.fromkeys([c for c in cells_text if re.match(r'^0\.\d+$',c)]))
                if nums: data['faktor'] = nums[-1]
            if 'Ploština korisne površine' in full and 'A K,' in full:
                nums = list(dict.fromkeys([c for c in cells_text if re.match(r'^\d{2,}\.\d+$',c)]))
                if nums: data['ak'] = nums[-1]
            if 'kondicionirane' in full and 'dimenzijama' in full:
                nums = list(dict.fromkeys([c for c in cells_text if re.match(r'^\d{2,}\.\d+$',c)]))
                if nums: data['bruto'] = nums[-1]
            if 'Ukupna ploština pročelja' in full:
                nums = list(dict.fromkeys([c for c in cells_text if re.match(r'^\d{2,}\.\d+$',c)]))
                if nums: data['procelj'] = nums[-1]
            if 'Ukupna ploština prozora' in full:
                nums = list(dict.fromkeys([c for c in cells_text if re.match(r'^\d+\.\d+$',c)]))
                if nums: data['prozori'] = nums[-1]

            # Systems
            if 'Vrsta dizalice topline' in full:
                pool = list(dict.fromkeys([c for c in cells_text if c and 'Vrsta' not in c and 'dizalice' not in c and c!='-']))
                if pool: data['grijIzvor'] = 'Dizalica topline ' + pool[-1]
            if 'Učinak u definiranoj radnoj točki' in full:
                nums = list(dict.fromkeys([c for c in cells_text if re.match(r'^\d+\.\d+$',c)]))
                if nums: data['grijSnaga'] = nums[-1]
            if 'Sezonski toplinski množitelj' in full and 'grijanj' in full:
                floats = list(dict.fromkeys([c for c in cells_text if re.match(r'^\d+\.\d+$',c)]))
                if floats: data['grijCop'] = 'SCOP = ' + floats[-1]
            if 'Direktno grijani električni' in full and any('DA'==c for c in cells_text):
                data['ptvTip'] = 'Direktno grijani električni spremnik (DGA)'
            if 'Volumen' in full and ('spremnik' in full.lower() or 'akumul' in full.lower()):
                nums = list(dict.fromkeys([c for c in cells_text if re.match(r'^\d+\.?\d*$',c)]))
                if nums: data['ptvVol'] = nums[-1]
            if 'Nazivna snaga' in full and ('grijač' in full.lower() or 'električ' in full.lower()):
                nums = list(dict.fromkeys([c for c in cells_text if re.match(r'^\d+\.?\d*$',c)]))
                if nums: data['ptvSnaga'] = nums[-1]
            if 'Nema definiranih sustava hlađenja' in full:
                data['hladVrsta'] = 'Nema – sustav hlađenja nije definiran'
            if 'Godišnja potrebna toplinska energija za PTV' in full:
                nums = list(dict.fromkeys([c for c in cells_text if re.match(r'^\d+\.\d+$',c)]))
                if nums: data['ptvQw'] = nums[-1]

            # Specific climate data from table 9 (2.A.5.4 Rezultati proračuna)
            if 'Q H,nd' in full and 'spec' in full.lower():
                nums = list(dict.fromkeys([c for c in cells_text if re.match(r'^\d+\.\d+$',c)]))
                if nums: data['qhndSpec'] = nums[-1]
            if 'E prim' in full and 'spec' in full.lower() and 'jedinici' in full:
                nums = list(dict.fromkeys([c for c in cells_text if re.match(r'^\d+\.\d+$',c)]))
                if nums: data['eprimSpec'] = nums[-1]
            if 'E del' in full and 'spec' in full.lower():
                nums = list(dict.fromkeys([c for c in cells_text if re.match(r'^\d+\.\d+$',c)]))
                if nums: data['edelSpec'] = nums[-1]
            if 'OIE' in full and 'spec' in full.lower():
                nums = list(dict.fromkeys([c for c in cells_text if re.match(r'^\d+\.\d+$',c)]))
                if nums: data['oieSpec'] = nums[-1]

            # U-values (real cells from lxml: [naziv, A, U, Umax, OK])
            non_empty = [c for c in cells_text if c and c not in ['-','']]
            if len(non_empty) < 3: continue
            naziv = non_empty[0]
            if any(naziv.startswith(s) for s in SKIP): continue
            if not any(v in naziv for v in VALID): continue
            if naziv in seen_names: continue
            rest = non_empty[1:]
            nums = [c for c in rest if re.match(r'^\d+\.?\d*$',c)]
            if len(nums) < 2: continue
            if len(nums) >= 3:
                area_val=nums[0]; u_str=nums[1]; umax_str=nums[2]
            else:
                area_val='—'; u_str=nums[0]; umax_str=nums[1]
            u_val=sf(u_str); umax_val=sf(umax_str)
            if u_val is None or umax_val is None or umax_val==0: continue
            seen_names.add(naziv)
            uvals.append({'naziv':naziv,'area':area_val,'u':u_str,'umax':umax_str,
                          'provjera':'ZADOVOLJAVA' if u_val<=umax_val else 'NE ZADOVOLJAVA'})

        data['uvalues'] = uvals
        data['kiRefRaw'] = '\n'.join(raw_lines)

    # Energy classes
    def cq(v):
        v=sf(v)
        if v is None: return '—'
        if v<=15: return 'A+'
        if v<=30: return 'A'
        if v<=50: return 'B'
        if v<=75: return 'C'
        if v<=100: return 'D'
        if v<=150: return 'E'
        if v<=200: return 'F'
        return 'G'
    def ce(v,mx):
        v=sf(v); mx=sf(mx)
        if not v or not mx: return '—'
        r=v/mx
        if r<=0.25: return 'A+'
        if r<=0.50: return 'A'
        if r<=0.75: return 'B'
        if r<=1.00: return 'C'
        if r<=1.50: return 'D'
        if r<=2.00: return 'E'
        if r<=2.50: return 'F'
        return 'G'
    data['razredQhnd'] = cq(data.get('qhndM2'))
    data['razredEprim'] = ce(data.get('eprimM2'), data.get('eprimMax'))
    return data

if __name__ == '__main__':
    if len(sys.argv)<2:
        print(json.dumps({'error':'No file path'})); sys.exit(1)
    try:
        print(json.dumps(extract(sys.argv[1]), ensure_ascii=False))
    except Exception as e:
        import traceback
        print(json.dumps({'error':str(e),'trace':traceback.format_exc()[-800:]}))
        sys.exit(1)
